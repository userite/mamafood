#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""Convert DOCX to HTML and extract formula (iskustveno mlqko) data"""

try:
    from docx import Document
    from bs4 import BeautifulSoup
    import html
    
    def convert_docx_to_html(docx_file):
        """Convert DOCX to HTML"""
        doc = Document(docx_file)
        
        html_content = "<html><head><meta charset='UTF-8'></head><body>"
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            html_content += f"<h2>Table {i+1}</h2>"
            html_content += "<table border='1'>"
            
            for row in table.rows:
                html_content += "<tr>"
                for cell in row.cells:
                    html_content += f"<td>{cell.text}</td>"
                html_content += "</tr>"
            
            html_content += "</table>"
        
        # Extract paragraphs
        html_content += "<h2>Text Content</h2>"
        for para in doc.paragraphs:
            if para.text.strip():
                html_content += f"<p>{para.text}</p>"
        
        html_content += "</body></html>"
        
        # Parse and clean
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup.prettify()
    
    # Main
    docx_file = 'mamafood_formula_full.docx'
    
    try:
        html_output = convert_docx_to_html(docx_file)
        
        # Save
        with open('mamafood_formula_full.html', 'w', encoding='utf-8') as f:
            f.write(html_output)
        
        print("OK Converted successfully to mamafood_formula_full.html")
        print("\nContent preview (first 2000 chars):")
        print(html_output[:2000])
        
    except Exception as e:
        print(f"Error: {e}")
        print("\nTrying alternative method...")
        
        # Alternative: Read as binary and extract
        import zipfile
        import xml.etree.ElementTree as ET
        
        with zipfile.ZipFile(docx_file, 'r') as docx_zip:
            # Get document.xml
            xml_content = docx_zip.read('word/document.xml')
            root = ET.fromstring(xml_content)
            
            html_output = "<html><head><meta charset='UTF-8'></head><body>"
            html_output += "<h1>Formula (Изкуствено мляко) Data</h1>"
            
            # Extract text
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    html_output += f"<p>{elem.text}</p>"
            
            html_output += "</body></html>"
            
            with open('mamafood_formula_full.html', 'w', encoding='utf-8') as f:
                f.write(html_output)
            
            print("OK Converted successfully (alternative method)")
            print("\nContent preview (first 2000 chars):")
            print(html_output[:2000])

except ImportError as e:
    print(f"Missing library: {e}")
    print("Install with: pip install python-docx beautifulsoup4")
    print("Or use built-in libraries...")
    
    # Use built-in zipfile and xml
    import zipfile
    import xml.etree.ElementTree as ET
    from html import escape as html_escape
    
    docx_file = 'mamafood_formula_full.docx'
    
    try:
        with zipfile.ZipFile(docx_file, 'r') as docx_zip:
            # Get document.xml
            xml_content = docx_zip.read('word/document.xml')
            root = ET.fromstring(xml_content)
            
            html_output = "<html><head><meta charset='UTF-8'></head><body>"
            html_output += "<h1>Formula (Изкуствено мляко) - Данни</h1>"
            
            # Extract text from all elements
            text_content = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    text_content.append(elem.text.strip())
            
            html_output += "<pre>"
            for text in text_content:
                if text:
                    html_output += f"{html_escape(text)}\n"
            html_output += "</pre>"
            
            html_output += "</body></html>"
            
            with open('mamafood_formula_full.html', 'w', encoding='utf-8') as f:
                f.write(html_output)
            
            print("OK Converted successfully!")
            # Skip console output to avoid encoding issues
            print(f"\nFile created: mamafood_formula_full.html")
            
    except Exception as e:
        print(f"Error: {e}")

