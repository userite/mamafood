#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""Convert КРАТКО_РЪКОВОДСТВО.md to .docx format"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    def parse_markdown_to_docx(md_file, docx_file):
        """Convert markdown file to docx"""
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Headers
            if line.startswith('# '):
                # H1
                text = line[2:].strip()
                p = doc.add_heading(text, level=1)
                i += 1
            elif line.startswith('## '):
                # H2
                text = line[3:].strip()
                p = doc.add_heading(text, level=2)
                i += 1
            elif line.startswith('### '):
                # H3
                text = line[4:].strip()
                p = doc.add_heading(text, level=3)
                i += 1
            elif line.startswith('#### '):
                # H4
                text = line[5:].strip()
                p = doc.add_heading(text, level=4)
                i += 1
            elif line.startswith('---'):
                # Horizontal rule - add spacing
                doc.add_paragraph()
                i += 1
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet list
                text = line[2:].strip()
                # Remove markdown bold **text**
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                # Remove markdown code `text`
                text = re.sub(r'`(.*?)`', r'\1', text)
                p = doc.add_paragraph(text, style='List Bullet')
                i += 1
            elif line.startswith('1. ') or re.match(r'^\d+\.\s', line):
                # Numbered list
                text = re.sub(r'^\d+\.\s', '', line)
                # Remove markdown bold **text**
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                # Remove markdown code `text`
                text = re.sub(r'`(.*?)`', r'\1', text)
                p = doc.add_paragraph(text, style='List Number')
                i += 1
            else:
                # Regular paragraph
                text = line
                # Remove markdown bold **text** but keep the text
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                # Remove markdown code `text`
                text = re.sub(r'`(.*?)`', r'\1', text)
                # Remove markdown links [text](url) -> text
                text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
                
                if text.strip():
                    p = doc.add_paragraph(text)
                i += 1
        
        # Save document
        doc.save(docx_file)
        import sys
        sys.stdout.reconfigure(encoding='utf-8')
        print(f"Успешно създаден: {docx_file}")
        
    # Main
    md_file = 'КРАТКО_РЪКОВОДСТВО.md'
    docx_file = 'КРАТКО_РЪКОВОДСТВО.docx'
    
    parse_markdown_to_docx(md_file, docx_file)
    
except ImportError as e:
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    print(f"Липсва библиотека: {e}")
    print("\nИнсталирай с:")
    print("pip install python-docx")
    print("\nИли на Windows:")
    print("python -m pip install python-docx")
    
except FileNotFoundError as e:
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    print(f"Файлът не е намерен: {e}")
    print("Уверете се, че файлът 'КРАТКО_РЪКОВОДСТВО.md' съществува в същата директория.")
    
except Exception as e:
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    print(f"Грешка: {e}")
    import traceback
    traceback.print_exc()

